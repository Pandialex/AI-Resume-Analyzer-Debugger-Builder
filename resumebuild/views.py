# resumebuild/views.py

import os
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse, Http404
from django.template.loader import render_to_string  # <-- Correctly imported for the preview
from .models import ResumeTemplate

# --- Imports for DOCX generation ---
from docx import Document
from io import BytesIO

def build_resume_view(request):
    """
    Renders the main builder page with all available templates.
    This is the first page the user sees.
    """
    templates = ResumeTemplate.objects.filter(is_active=True)
    return render(request, 'resumebuild/build_form.html', {'templates': templates})

def preview_resume_view(request, template_id):
    """
    Renders the resume HTML for the live preview iframe.
    It takes the template_id and all form data from the URL parameters.
    """
    template = get_object_or_404(ResumeTemplate, pk=template_id, is_active=True)
    
    # Get all data from the URL parameters (e.g., ?full_name=John&email=...)
    data = request.GET.dict()

    # Render the template's HTML file with the provided data
    try:
        html_content = render_to_string(template.html_file.path, {'data': data})
        return HttpResponse(html_content)
    except FileNotFoundError:
        # This error happens if the HTML file path in the model is incorrect
        raise Http404("Template HTML file not found.")

def download_resume_view(request, template_id, format):
    """
    Generates and returns a downloadable DOCX file.
    The PDF option has been removed to avoid dependency issues.
    """
    template = get_object_or_404(ResumeTemplate, pk=template_id, is_active=True)
    data = request.GET.dict()
    user_name = data.get('full_name', 'resume').replace(' ', '_')

    # --- SIMPLIFIED: Only handle DOCX now ---
    if format.lower() == 'docx':
        return _generate_docx(data, user_name)
    else:
        # This case will be hit if someone tries to access the PDF URL directly
        raise Http404("PDF download is not available. Please use the DOCX download.")

def _generate_docx(data, user_name):
    """
    Helper function to generate a DOCX file using the python-docx library.
    It takes the form data and structures it into a Word document.
    """
    doc = Document()
    
    # --- Title and Contact ---
    doc.add_heading(data.get('full_name', 'Your Name'), 0)
    p = doc.add_paragraph()
    p.add_run(f"📧 {data.get('email', '')} | 📱 {data.get('phone', '')} | 📍 {data.get('address', '')}")
    
    # --- Job Title ---
    if data.get('job_title'):
        doc.add_heading(data.get('job_title'), level=1)

    # --- Profile Summary ---
    if data.get('summary'):
        doc.add_heading('Professional Summary', level=2)
        doc.add_paragraph(data.get('summary'))

    # --- Experience ---
    if data.get('experience'):
        doc.add_heading('Experience', level=2)
        # Simple split by newline for multiple experiences
        for exp in data.get('experience').split('\n'):
            if exp.strip():
                doc.add_paragraph(exp.strip(), style='List Bullet')

    # --- Projects (for Freshers) ---
    if data.get('projects'):
        doc.add_heading('Projects', level=2)
        for proj in data.get('projects').split('\n'):
            if proj.strip():
                doc.add_paragraph(proj.strip(), style='List Bullet')

    # --- Education ---
    if data.get('education'):
        doc.add_heading('Education', level=2)
        for edu in data.get('education').split('\n'):
            if edu.strip():
                doc.add_paragraph(edu.strip(), style='List Bullet')
                
    # --- Certifications (for Freshers) ---
    if data.get('certifications'):
        doc.add_heading('Certifications', level=2)
        for cert in data.get('certifications').split('\n'):
            if cert.strip():
                doc.add_paragraph(cert.strip(), style='List Bullet')

    # --- Skills ---
    if data.get('skills'):
        doc.add_heading('Skills', level=2)
        doc.add_paragraph(data.get('skills'))

    # --- Save the document to an in-memory buffer ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # --- Create the HttpResponse to trigger the download ---
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{user_name}_resume.docx"'
    return response
